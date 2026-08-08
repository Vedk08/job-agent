"""Takes the user's original CV (as DOCX bytes) plus a set of accepted bullet
changes, and produces a new CV that keeps the *exact* original template —
fonts, spacing, bullet styling, header layout, everything — because we only
ever rewrite the text inside existing paragraphs. We never rebuild the
document from scratch.

Why this approach instead of generating a fresh document:
  The user's requirement is "must match my current template exactly." The
  only way to guarantee that with python-docx is to start from their real
  .docx and mutate specific paragraphs' text runs — not describe a layout
  and hope it renders the same.

Output is DOCX only. A DOCX -> PDF path via headless LibreOffice was
prototyped and works, but was deliberately dropped: it roughly doubles the
Docker image size and adds real RAM pressure on Render's free 512MB tier.
DOCX-only keeps the deployed app dependency-free at the system level — the
user can export to PDF themselves from Word/Google Docs when needed.
"""
from dataclasses import dataclass
from io import BytesIO

import docx


@dataclass
class BulletChange:
    paragraph_index: int  # must match a Bullet.paragraph_index from cv_parser
    new_text: str


class CVGenerationError(Exception):
    pass


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    """Replace a paragraph's visible text while preserving its formatting.

    Strategy: keep the paragraph's FIRST run (which carries the font/size/
    bold/etc. that Word actually renders) and just swap its text. Any
    additional runs are cleared to empty — rather than deleted outright,
    since removing run XML elements mid-iteration is fiddly with
    python-docx — so no duplicate or stale text can leak through.
    """
    if not paragraph.runs:
        # Bullet paragraph got created with a run-less state — set text via
        # add_run so it still inherits the paragraph's own style/formatting.
        paragraph.add_run(new_text)
        return

    first_run = paragraph.runs[0]
    first_run.text = new_text
    for extra_run in paragraph.runs[1:]:
        extra_run.text = ""


def apply_changes(original_docx_bytes: bytes, changes: list[BulletChange]) -> bytes:
    """Returns new DOCX bytes with the accepted bullet changes applied."""
    document = docx.Document(BytesIO(original_docx_bytes))
    paragraphs = document.paragraphs

    by_index = {c.paragraph_index: c.new_text for c in changes}

    applied = 0
    for idx, new_text in by_index.items():
        if idx < 0 or idx >= len(paragraphs):
            raise CVGenerationError(
                f"paragraph_index {idx} is out of range for this document "
                f"({len(paragraphs)} paragraphs) — the CV may have changed "
                f"since it was parsed; re-upload and re-analyze."
            )
        _replace_paragraph_text(paragraphs[idx], new_text)
        applied += 1

    if applied != len(changes):
        raise CVGenerationError("Some bullet changes could not be applied.")

    out = BytesIO()
    document.save(out)
    return out.getvalue()


def generate_cv(original_docx_bytes: bytes, changes: list[BulletChange]) -> dict:
    """End-to-end: apply changes, return the new DOCX.

    Returns {"docx": bytes}. Kept as a dict (rather than raw bytes) so the
    API layer's contract doesn't need to change if a PDF path is ever added
    back later.
    """
    new_docx = apply_changes(original_docx_bytes, changes)
    return {"docx": new_docx}
