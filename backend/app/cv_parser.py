"""Parses an uploaded CV (DOCX or PDF) into:
  1. plain text (fed to the LLM for analysis)
  2. a structured bullet map (only for DOCX — used later to edit the file in
     place without touching layout/formatting)

Design note: we deliberately do NOT try to fully "understand" the CV's
sections generically. We identify bullets by python-docx paragraph style
("List Paragraph") and record their paragraph index. That index is a stable
handle we can hand back to cv_generator.py later to edit just that one
paragraph's text, leaving every other paragraph (and all run-level
formatting: fonts, bold, spacing) completely untouched.

PDFs are read-only inputs here — if someone uploads a PDF CV, we can extract
text to analyze it, but we can't safely edit a PDF in place and preserve
layout, so CV *generation* only works from an uploaded DOCX (see
cv_generator.py). The API should surface that constraint to the frontend.
"""
from dataclasses import dataclass
from io import BytesIO
from typing import Optional

import docx
import pdfplumber


@dataclass
class Bullet:
    paragraph_index: int  # index into document.paragraphs — stable edit handle
    text: str
    # Best-effort context so the LLM knows what section/role this bullet
    # belongs to when suggesting a rewrite (e.g. "Cinematic Scene Understanding")
    section_heading: Optional[str]


@dataclass
class ParsedCV:
    plain_text: str
    bullets: list[Bullet]  # empty for PDF input
    source_format: str  # "docx" | "pdf"


def _is_heading_candidate(paragraph) -> bool:
    """Heuristic: short, non-bulleted paragraphs are section/entry headings.

    This CV format doesn't use Word heading styles consistently, so we infer
    headings from shape: not a List Paragraph, reasonably short, not empty.
    """
    if paragraph.style and paragraph.style.name == "List Paragraph":
        return False
    text = paragraph.text.strip()
    return bool(text) and len(text) < 90


def parse_docx(file_bytes: bytes) -> ParsedCV:
    document = docx.Document(BytesIO(file_bytes))

    plain_lines: list[str] = []
    bullets: list[Bullet] = []
    # Keep the last few heading-like lines seen. The *title* (project name,
    # "Company, City", role name) is usually 2-3 lines above the first
    # bullet, with tech-stack/link lines in between — so we join a short
    # window rather than trusting only the immediately preceding line.
    recent_headings: list[str] = []
    MAX_HEADING_WINDOW = 3

    for idx, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue

        plain_lines.append(text)

        is_bullet = paragraph.style and paragraph.style.name == "List Paragraph"
        if is_bullet:
            bullets.append(Bullet(
                paragraph_index=idx,
                text=text,
                section_heading=" | ".join(recent_headings) or None,
            ))
        elif _is_heading_candidate(paragraph):
            recent_headings.append(text)
            recent_headings = recent_headings[-MAX_HEADING_WINDOW:]

    return ParsedCV(
        plain_text="\n".join(plain_lines),
        bullets=bullets,
        source_format="docx",
    )


def parse_pdf(file_bytes: bytes) -> ParsedCV:
    text_parts: list[str] = []
    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(page_text)

    return ParsedCV(
        plain_text="\n".join(text_parts),
        bullets=[],  # can't safely edit a PDF's layout in place
        source_format="pdf",
    )


def parse_cv(filename: str, file_bytes: bytes) -> ParsedCV:
    lower = filename.lower()
    if lower.endswith(".docx"):
        return parse_docx(file_bytes)
    if lower.endswith(".pdf"):
        return parse_pdf(file_bytes)
    raise ValueError(f"Unsupported CV file type: {filename} (accepts .docx or .pdf)")
